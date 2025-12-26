"""Document Loader - Extract text from various document formats."""

import logging
from typing import Dict, Any, Optional
from pathlib import Path
import mimetypes

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Load and extract text from various document formats."""
    
    SUPPORTED_FORMATS = {
        '.txt': 'text/plain',
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.md': 'text/markdown'
    }
    
    @staticmethod
    async def load(file_path: str) -> Dict[str, Any]:
        """
        Load document and extract text.
        
        Args:
            file_path: Path to document
            
        Returns:
            Dict with text and metadata
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        suffix = path.suffix.lower()
        
        if suffix == '.txt' or suffix == '.md':
            return await DocumentLoader._load_text(path)
        elif suffix == '.pdf':
            return await DocumentLoader._load_pdf(path)
        elif suffix == '.docx':
            return await DocumentLoader._load_docx(path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
    
    @staticmethod
    async def _load_text(path: Path) -> Dict[str, Any]:
        """Load plain text file."""
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        return {
            'text': text,
            'metadata': {
                'filename': path.name,
                'format': 'text',
                'size': path.stat().st_size
            }
        }
    
    @staticmethod
    async def _load_pdf(path: Path) -> Dict[str, Any]:
        """Load PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not installed. Install with: pip install PyPDF2")
        
        text_parts = []
        
        with open(path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            num_pages = len(pdf_reader.pages)
            
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text_parts.append(page.extract_text())
        
        text = '\n\n'.join(text_parts)
        
        return {
            'text': text,
            'metadata': {
                'filename': path.name,
                'format': 'pdf',
                'size': path.stat().st_size,
                'pages': num_pages
            }
        }
    
    @staticmethod
    async def _load_docx(path: Path) -> Dict[str, Any]:
        """Load DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        doc = DocxDocument(path)
        text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        text = '\n\n'.join(text_parts)
        
        return {
            'text': text,
            'metadata': {
                'filename': path.name,
                'format': 'docx',
                'size': path.stat().st_size,
                'paragraphs': len(doc.paragraphs)
            }
        }
